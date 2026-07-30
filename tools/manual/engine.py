# -*- coding: utf-8 -*-
"""Движок сборки руководства пользователя iwo Helper Desktop (.docx).

Здесь только ОФОРМЛЕНИЕ и строительные блоки: стили, поля Word, сноски, рисунки,
таблицы. Текст руководства лежит отдельно, по модулю на язык (body_ru, body_en), и
собирается через те же самые функции — так два руководства не могут разойтись
устройством, а разойтись словами им и положено.

Прежнее назначение:

Оформление задано заказчиком: Times New Roman везде, заголовки 16 пт, текст 14 пт,
красная строка, полуторный интервал, рисунки подписаны «Рис. N» и упомянуты в тексте,
нумерация страниц, автоматическое оглавление, сноски, титульный лист.

Снимки экрана — настоящие окна программы (см. Shots.cs), диаграмма построена по
измеренным размерам файлов (см. measure_compression.py). Ничего не нарисовано «от руки».
"""
import io
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = None  # задаётся ниже, когда известен язык
REPO = r"C:\work\iwoHelperDesktop"
OUT_DIR = os.path.join(REPO, "docs")

# Язык и всё, что от него зависит: имя файла и служебные подписи документа.
LANG = os.environ.get("MANUAL_LANG", "ru")
LABELS = {
    "ru": {
        "file": "Инструкция пользователя.docx",
        "shots": "shots-ru",
        "figRef": "(рис. %d)",
        "figCaption": "Рис. %d. %s",
        "tableCaption": "Таблица %d — %s",
        "tableRef": "табл. %d",
        "note": "Примечание. ",
        "toc": "ОГЛАВЛЕНИЕ",
        "titleWord": "РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ",
        "subtitle": "Настольные инструменты для работы с документами Excel и PDF",
        "version": "Версия программы – %s",
        "tocHint": "Оглавление собирается автоматически. Если оно не заполнилось, "
                   "выделите документ клавишами Ctrl+A и нажмите F9.",
        "chartTitle": "Размер файла, МБ",
        "chartLevels": ["Отлично\n(без сжатия)", "Хорошо\n(150 dpi)", "Нормально\n(72 dpi)"],
        "chartValue": "%.2f МБ",
        "decimal": ",",
    },
    "en": {
        "file": "User manual.docx",
        "shots": "shots-en",
        "figRef": "(fig. %d)",
        "figCaption": "Fig. %d. %s",
        "tableCaption": "Table %d — %s",
        "tableRef": "table %d",
        "note": "Note. ",
        "toc": "CONTENTS",
        "titleWord": "USER MANUAL",
        "subtitle": "Desktop tools for working with Excel and PDF documents",
        "version": "Program version – %s",
        "tocHint": "The contents are collected automatically. If they came out empty, "
                   "select the document with Ctrl+A and press F9.",
        "chartTitle": "File size, MB",
        "chartLevels": ["Excellent\n(no compression)", "Good\n(150 dpi)", "Normal\n(72 dpi)"],
        "chartValue": "%.2f MB",
        "decimal": ".",
    },
}
L = LABELS[LANG]
SHOTS = os.path.join(HERE, L["shots"])
OUT = os.path.join(OUT_DIR, L["file"])


def app_version():
    """Версия программы — из AssemblyInfo.cs, единственного её источника в репозитории.

    Вписанный руками номер молча устаревает: инструкция вшита ресурсом в exe, и
    руководство с чужой версией на титуле уезжает в релиз незамеченным.
    """
    import re
    path = os.path.join(REPO, "src", "AssemblyInfo.cs")
    with io.open(path, encoding="utf-8-sig") as f:
        match = re.search(r'AssemblyInformationalVersion\("([^"]+)"\)', f.read())
    if not match:
        raise SystemExit("в AssemblyInfo.cs не найдена AssemblyInformationalVersion")
    return match.group(1)


VERSION = app_version()

BODY_PT = Pt(14)
HEAD_PT = Pt(16)
NOTE_PT = Pt(12)          # сноски — служебный аппарат, набираются мельче основного текста
FONT = "Times New Roman"
INDENT = Cm(1.25)         # красная строка

# ---------------------------------------------------------------- диаграмма

def build_chart():
    """Диаграмма размеров файла при разных уровнях сжатия — по измеренным значениям."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    data = {}
    for line in io.open(os.path.join(HERE, "compression.txt"), encoding="utf-8"):
        key, value = line.strip().split("=")
        data[key] = int(value) / 1048576.0

    labels = ["Отлично\n(без сжатия)", "Хорошо\n(150 dpi)", "Нормально\n(72 dpi)"]
    values = [data["none"], data["good"], data["small"]]
    colors = ["#8c8f94", "#0f6cbd", "#107c41"]

    plt.rcParams["font.family"] = FONT
    fig, ax = plt.subplots(figsize=(7.4, 3.5), dpi=200)
    bars = ax.bar(labels, values, color=colors, width=0.55)
    top = max(values)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + top * 0.035,
                (L["chartValue"] % value).replace(".", L["decimal"]), ha="center", va="bottom", fontsize=12)
        share = 100.0 * value / values[0]
        # Белым внутри столбика — только если столбик достаточно высок, иначе подпись
        # сольётся с фоном. Низкому столбику подпись ставим над ним тёмной.
        if value > top * 0.18:
            ax.text(bar.get_x() + bar.get_width() / 2, value / 2,
                    "%.0f %%" % share, ha="center", va="center",
                    fontsize=13, color="white", fontweight="bold")
        else:
            ax.text(bar.get_x() + bar.get_width() / 2, value + top * 0.12,
                    "%.0f %%" % share, ha="center", va="bottom",
                    fontsize=13, color=bar.get_facecolor(), fontweight="bold")
    ax.set_ylabel(L["chartTitle"], fontsize=12)
    ax.set_ylim(0, top * 1.22)
    ax.tick_params(labelsize=12)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    path = os.path.join(SHOTS, "chart.png")
    fig.savefig(path)
    plt.close(fig)
    return values


CHART = build_chart()

# ---------------------------------------------------------------- документ

doc = Document()


def set_font(run, size=BODY_PT, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), FONT)


def style_document():
    """Единое оформление: шрифт, кегли, красная строка, полуторный интервал, поля."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_PT
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), FONT)
    fmt = normal.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = INDENT
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Заголовки — встроенными стилями: только их подхватывает поле автооглавления.
    for name, before, after in (("Heading 1", 18, 10), ("Heading 2", 14, 8), ("Heading 3", 12, 6)):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = HEAD_PT
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        heading_fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        # Ссылка на шрифт темы сильнее явного имени: пока эти атрибуты на месте,
        # заголовок остаётся рубленым, чем бы мы его ни назначали.
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if heading_fonts.get(qn(attr)) is not None:
                del heading_fonts.attrib[qn(attr)]
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            heading_fonts.set(qn(attr), FONT)
        hf = style.paragraph_format
        hf.first_line_indent = Cm(0)
        # Заголовки — по центру: так принято в русских руководствах, и так их
        # выключил заказчик. Выключка задана В СТИЛЕ, а не на каждом абзаце:
        # заголовков четыре десятка, и прямое форматирование пришлось бы
        # повторять в h1/h2/h3 и не забыть при добавлении нового уровня.
        hf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        hf.space_before = Pt(before)
        hf.space_after = Pt(after)
        hf.keep_with_next = True

    toc_styles()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)



def toc_styles():
    """Объявить стили строк оглавления заранее.

    Оглавление собирает сам Word, и если стилей «toc 1»/«toc 2» в документе нет, он
    создаёт их по шаблону — со шрифтом темы. Тогда в готовом файле появляется Cambria,
    хотя мы её нигде не просили.
    """
    from docx.enum.style import WD_STYLE_TYPE
    for level, name in ((1, "toc 1"), (2, "toc 2"), (3, "toc 3")):
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = BODY_PT
        fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if fonts.get(qn(attr)) is not None:
                del fonts.attrib[qn(attr)]
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            fonts.set(qn(attr), FONT)
        fmt = style.paragraph_format
        fmt.first_line_indent = Cm(0)
        fmt.left_indent = Cm(0.75 * (level - 1))
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.space_after = Pt(0)

def field(paragraph, instruction, placeholder=""):
    """Поле Word (PAGE, TOC): значение считает сам Word, а не мы."""
    begin = paragraph.add_run()
    element = OxmlElement("w:fldChar")
    element.set(qn("w:fldCharType"), "begin")
    element.set(qn("w:dirty"), "true")
    begin._element.append(element)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._element.append(instr)

    sep = paragraph.add_run()
    element = OxmlElement("w:fldChar")
    element.set(qn("w:fldCharType"), "separate")
    sep._element.append(element)

    if placeholder:
        set_font(paragraph.add_run(placeholder))

    end = paragraph.add_run()
    element = OxmlElement("w:fldChar")
    element.set(qn("w:fldCharType"), "end")
    end._element.append(element)
    for run in (begin, instr_run, sep, end):
        set_font(run)


def page_numbers():
    """Номера страниц по центру нижнего колонтитула; на титульном листе номера нет."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_footer.is_linked_to_previous = False
    for paragraph in section.first_page_footer.paragraphs:
        paragraph.text = ""
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Cm(0)
    field(footer, " PAGE ", "2")


def update_fields_on_open():
    """Попросить Word пересчитать поля при открытии — иначе оглавление будет пустым."""
    settings = doc.settings.element
    flag = OxmlElement("w:updateFields")
    flag.set(qn("w:val"), "true")
    settings.append(flag)


# ---------------------------------------------------------------- сноски

FOOTNOTES = []
FOOTNOTE_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)


def footnote(paragraph, text):
    """Обычная сноска внизу страницы. python-docx их не умеет, собираем часть пакета сами."""
    number = len(FOOTNOTES) + 1
    FOOTNOTES.append(text)
    run = paragraph.add_run()
    set_font(run)
    rpr = run._element.get_or_add_rPr()
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rpr.append(vert)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(number))
    run._element.append(ref)
    return paragraph


def attach_footnotes():
    if not FOOTNOTES:
        return
    parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
             "<w:footnotes %s>" % FOOTNOTE_NS,
             '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" '
             'w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>',
             '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" '
             'w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>']
    size = str(int(NOTE_PT.pt * 2))
    for index, text in enumerate(FOOTNOTES, start=1):
        safe = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        rpr = ('<w:rPr><w:rFonts w:ascii="%s" w:hAnsi="%s" w:cs="%s" w:eastAsia="%s"/>'
               '<w:sz w:val="%s"/><w:szCs w:val="%s"/></w:rPr>' % (FONT, FONT, FONT, FONT, size, size))
        parts.append(
            '<w:footnote w:id="%d"><w:p><w:pPr><w:spacing w:after="0" w:line="240" '
            'w:lineRule="auto"/><w:ind w:firstLine="0"/><w:jc w:val="both"/></w:pPr>'
            '<w:r>%s<w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteRef/></w:r>'
            '<w:r>%s<w:t xml:space="preserve"> %s</w:t></w:r></w:p></w:footnote>'
            % (index, rpr, rpr, safe))
    parts.append("</w:footnotes>")
    blob = "".join(parts).encode("utf-8")

    partname = PackURI("/word/footnotes.xml")
    content_type = ("application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.footnotes+xml")
    part = Part(partname, content_type, blob, doc.part.package)
    doc.part.relate_to(part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")


# ---------------------------------------------------------------- строительные блоки


def retheme(path):
    """Переписать шрифты темы в сохранённом файле: Word берёт их для того, что создаёт сам."""
    import re
    import shutil
    import zipfile

    temp = path + ".tmp"
    with zipfile.ZipFile(path) as src, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.startswith("word/theme/"):
                xml = data.decode("utf-8")
                xml = re.sub(r'(<a:latin[^/>]*typeface=")[^"]*(")', r"\1" + FONT + r"\2", xml)
                xml = re.sub(r'(<a:ea[^/>]*typeface=")[^"]*(")', r"\1" + FONT + r"\2", xml)
                xml = re.sub(r'(<a:cs[^/>]*typeface=")[^"]*(")', r"\1" + FONT + r"\2", xml)
                data = xml.encode("utf-8")
            dst.writestr(item, data)
    shutil.move(temp, path)


def finalize_toc(path):
    """Собрать оглавление внутри .docx руками Word.

    python-docx умеет записать только ПОЛЕ оглавления: пункты и номера страниц
    считает Word, больше некому. Пока этого шага не было, в репозиторий уезжал
    файл с пустым оглавлением и просьбой нажать F9 — то есть работа, которую
    читатель должен был доделать за нас.

    Word водится из PowerShell (pywin32 в окружении нет), и путь ему передаётся
    ТОЛЬКО ASCII: Windows PowerShell 5.1 портит не-ASCII, а имя файла у нас
    кириллическое. Поэтому работаем по копии со служебным именем и возвращаем её
    на место средствами Python.
    """
    import shutil
    import subprocess

    ascii_copy = os.path.join(HERE, "_toc_finalize.docx")
    script = os.path.join(HERE, "finalize_toc.ps1")
    shutil.copyfile(path, ascii_copy)
    try:
        done = subprocess.run(
            ["powershell", "-NoProfile", "-File", script, "-Path", ascii_copy],
            capture_output=True, text=True)
        out = (done.stdout or "") + (done.stderr or "")
        # Молчаливый провал здесь означает пустое оглавление в релизе, поэтому
        # проверяем и код возврата, и метку успеха: сам по себе нулевой код
        # ничего не доказывает, если Word не открылся.
        if done.returncode != 0 or "ok" not in out:
            raise SystemExit("оглавление собрать не удалось (код %s):\n%s" % (done.returncode, out.strip()))
        print("оглавление:", " ".join(line for line in out.split() if line.startswith(("pages=", "toc_entries="))))
        shutil.move(ascii_copy, path)
    finally:
        if os.path.exists(ascii_copy):
            os.remove(ascii_copy)

def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def p(text="", indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False,
      size=BODY_PT, space_before=0, space_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = INDENT if indent else Cm(0)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if text:
        set_font(paragraph.add_run(text), size=size, bold=bold, italic=italic)
    return paragraph


def rich(*chunks):
    """Абзац из кусочков (текст, жирный?) — чтобы выделять названия кнопок."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = INDENT
    for chunk in chunks:
        text, bold = (chunk, False) if isinstance(chunk, str) else chunk
        set_font(paragraph.add_run(text), bold=bold)
    return paragraph


def listed(marker, text, bold_head=None, bold_marker=False):
    """Пункт перечня: маркер, необязательный жирный зачин, текст.

    Отступа слева и висячей строки НЕТ намеренно: пункт начинается с красной
    строки и переносится к левому полю - так набирают перечни в русских
    документах. Маркер рисуем сами: встроенный стиль списка берёт его из шрифта
    Symbol, а весь документ набран одним шрифтом. Маркированный пункт отличается
    от нумерованного ТОЛЬКО маркером, поэтому оба собираются здесь: иначе правка
    отступов требовала бы двух одинаковых правок в двух местах.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = INDENT
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    set_font(paragraph.add_run(marker), bold=bold_marker)
    if bold_head:
        set_font(paragraph.add_run(bold_head), bold=True)
    set_font(paragraph.add_run(text))
    return paragraph


def bullet(text, bold_head=None):
    return listed("\u2014\u00a0", text, bold_head)  # тире и неразрывный пробел


def step(number, text, bold_head=None):
    return listed("%d. " % number, text, bold_head, bold_marker=True)


def note(text):
    """Примечание — отдельный абзац с явной пометкой."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = INDENT
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    set_font(paragraph.add_run(L["note"]), bold=True)
    set_font(paragraph.add_run(text), italic=True)
    return paragraph


# --- рисунки: номера назначены заранее, чтобы ссылка в тексте не разошлась с подписью

FIG_ORDER = [
    "hub", "hub-pdf", "hub-lang", "help-split", "shortcuts", "merge-ctx", "goto", "preview",
    "merge-compress", "chart", "settings", "stats", "merge", "merge-menu",
    "split", "split-modes", "split-ranges", "split-everyn", "split-bookmarks",
    "split-template", "ops", "ops-dpi", "metadata", "ocr", "pptx", "excel", "excel-menu", "about",
]
_inserted = []


def fig(name):
    """Номер рисунка по его имени — для ссылки в тексте."""
    return FIG_ORDER.index(name) + 1


def ref(name):
    return L["figRef"] % fig(name)


def picture(name, caption):
    """Рисунок с подписью «Рис. N. …» и проверкой, что порядок вставки не разъехался."""
    expected = FIG_ORDER[len(_inserted)]
    if expected != name:
        raise SystemExit("рисунок «%s» вставлен вместо «%s» — порядок нарушен" % (name, expected))
    _inserted.append(name)

    path = os.path.join(SHOTS, name + ".png")
    with Image.open(path) as img:
        px, py = img.size
    # Размер рисунка — по ЧИТАЕМОСТИ, а не по числу точек: подписи внутри окон набраны
    # одним кеглем, поэтому целимся в постоянные ~110 точек на дюйм на бумаге. Ширина не
    # шире полосы набора (16 см), высота — не выше 17 см, иначе рисунок занимает страницу
    # целиком и гонит перед собой пустоту.
    width_cm = min(16.0, max(7.0, px / 110.0 * 2.54))
    if width_cm * py / float(px) > 17.0:
        width_cm = 17.0 * px / float(py)

    holder = doc.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.paragraph_format.first_line_indent = Cm(0)
    holder.paragraph_format.space_before = Pt(10)
    holder.paragraph_format.space_after = Pt(4)
    holder.paragraph_format.keep_with_next = True
    holder.add_run().add_picture(path, width=Cm(width_cm))

    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature.paragraph_format.first_line_indent = Cm(0)
    signature.paragraph_format.space_after = Pt(12)
    set_font(signature.add_run(L["figCaption"] % (fig(name), caption)))
    return signature



def repeat_header(row):
    """Повторять шапку таблицы на каждой странице: длинная таблица иначе продолжается
    безымянными колонками."""
    properties = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    properties.append(flag)

TABLES = []


def table(caption, headers, rows, widths=None):
    """Таблица с подписью «Таблица N — …» над ней, как принято в русских документах."""
    TABLES.append(caption)
    title = doc.add_paragraph()
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    set_font(title.add_run(L["tableCaption"] % (len(TABLES), caption)))

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, text in enumerate(headers):
        cell = t.rows[0].cells[index]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(0)
        set_font(para.add_run(text), bold=True)
    repeat_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = ""
            para = cells[index].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # в узких колонках выключка по ширине рвёт строки
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(0)
            set_font(para.add_run(text))
    if widths:
        t.autofit = False  # иначе Word пересчитает колонки по содержимому и заданные ширины пропадут
        for index, width in enumerate(widths):
            for row in t.rows:
                row.cells[index].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def table_ref():
    return L["tableRef"] % (len(TABLES) + 1)



def front_matter():
    """Титульный лист и поле оглавления — одинаковы для обоих языков, слова из L."""
    # ---------- титульный лист ----------

    for _ in range(4):
        p()
    p("iwo Helper Desktop", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(28))
    p()
    p(L["titleWord"], indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
      bold=True, size=Pt(20))
    p()
    p(L["subtitle"],
      indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    # Версия и год прижаты к низу титульного листа, как принято в руководствах.
    for _ in range(17):
        p()
    p(L["version"] % VERSION, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p()
    p("2026", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------- оглавление ----------

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.first_line_indent = Cm(0)
    toc_title.paragraph_format.space_after = Pt(12)
    # Новую страницу начинает САМ заголовок, а не отдельный абзац с разрывом:
    # такой абзац занимает строку, и на заполненном титульном листе она уезжает на
    # вторую страницу, оставляя её пустой, а оглавление — на третьей.
    toc_title.paragraph_format.page_break_before = True
    set_font(toc_title.add_run(L["toc"]), size=HEAD_PT, bold=True)

    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    field(toc, r' TOC \o "1-2" \h \z \u ',
          L["tocHint"])
    doc.add_page_break()




# Оформление применяет сборщик (build_manual.py) — движок сам ничего не начинает.

