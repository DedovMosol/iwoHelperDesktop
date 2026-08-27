# -*- coding: utf-8 -*-
"""Проверить готовые RU/EN руководства и регрессию обновления оглавления.

Проверяется именно сохранённый DOCX: стили, поля, сноски, картинки, подписи и
совпадение текущего оглавления с Heading 1/Heading 2. С ключом ``--toc-regression``
скрипт также работает с временной копией, добавляет два новых заголовка, запускает
тот же Word-финализатор и доказывает, что оба пункта с номерами страниц попали в
пересобранное оглавление. Исходные руководства не меняются.
"""
import argparse
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
DOCS = os.path.join(REPO, "docs")
FONT = "Times New Roman"
MANUALS = {
    "ru": {
        "file": "Инструкция пользователя.docx",
        "figure": r"Рис\. (\d+)\.",
        "mention": r"рис\. (\d+)",
        "table": r"Таблица (\d+) —",
        "review_caption": "Результат сравнения: ранняя версия слева, поздняя справа",
        "probe_h1": "11. Проверка нового раздела",
        "probe_h2": "11.1. Проверка нового подраздела",
    },
    "en": {
        "file": "User manual.docx",
        "figure": r"Fig\. (\d+)\.",
        "mention": r"fig\. (\d+)",
        "table": r"Table (\d+) —",
        "review_caption": "The comparison result: earlier version left, later version right",
        "probe_h1": "11. New section update probe",
        "probe_h2": "11.1. New subsection update probe",
    },
}


def inherited(style, attribute):
    """Свойство шрифта с учётом наследования стилей."""
    while style is not None:
        value = getattr(style.font, attribute)
        if value is not None:
            return value
        style = style.base_style
    return None


def toc_entries(document):
    return [paragraph for paragraph in document.paragraphs
            if paragraph.style.name.lower() in ("toc 1", "toc 2")]


def toc_text(paragraph):
    return re.sub(r"\t\d+$", "", paragraph.text).strip()


def verify_toc_update(path, spec, check, verify_open_request=False):
    """На временной копии доказать добавление новых Heading 1/2 в полный TOC."""
    before = Document(path)
    old_entries = set(toc_text(paragraph) for paragraph in toc_entries(before))
    source_hash = hashlib.sha256(open(path, "rb").read()).hexdigest()
    probe_dir = tempfile.mkdtemp(prefix="iwo-manual-toc-")
    probe = os.path.join(probe_dir, "manual.docx")
    shutil.copyfile(path, probe)
    try:
        edited = Document(probe)
        edited.add_heading(spec["probe_h1"], level=1)
        edited.add_heading(spec["probe_h2"], level=2)
        edited.save(probe)
        if verify_open_request:
            # The document keeps the standards-based request, but desktop Word 16
            # does not reliably rebuild a TOC merely by opening it. Make that product
            # limitation executable so nobody later re-labels the request as a
            # guarantee; the deterministic regression below invokes Update explicitly.
            ascii_probe = os.path.join(HERE, "_toc_open_probe_%d.docx" % os.getpid())
            shutil.copyfile(probe, ascii_probe)
            try:
                script = (
                    '$w=New-Object -ComObject Word.Application; '
                    '$w.Visible=$false; $w.DisplayAlerts=0; try {'
                    '$d=$w.Documents.Open(\"%s\",$false,$false); '
                    '$found=$d.TablesOfContents.Item(1).Range.Text.Contains(\"%s\"); '
                    '$d.Close($false); Write-Output $found'
                    '} finally {$w.Quit(); '
                    '[Runtime.InteropServices.Marshal]::ReleaseComObject($w)|Out-Null}'
                    % (ascii_probe.replace("'", "''"), spec["probe_h1"].replace("'", "''")))
                opened = subprocess.run(
                    ["powershell", "-NoProfile", "-Command", script],
                    cwd=HERE, capture_output=True, text=True, timeout=120)
                check(opened.returncode == 0 and "False" in (opened.stdout or ""),
                      "desktop Word не принят за гарант автообновления при открытии",
                      "поведение открытия Word изменилось или проба не завершилась: %s"
                      % (((opened.stdout or "") + (opened.stderr or "")).strip()))
            finally:
                if os.path.exists(ascii_probe):
                    os.remove(ascii_probe)
        done = subprocess.run(
            [sys.executable, os.path.join(HERE, "update_toc.py"), probe],
            cwd=HERE, capture_output=True, text=True)
        output = ((done.stdout or "") + (done.stderr or "")).strip()
        check(done.returncode == 0, "временная копия пересобрана Word",
              "Word не пересобрал временную копию (код %d): %s" % (done.returncode, output))
        if done.returncode != 0:
            return

        updated = Document(probe)
        entries = toc_entries(updated)
        by_text = {toc_text(paragraph): paragraph.text for paragraph in entries}
        check(spec["probe_h1"] in by_text and spec["probe_h2"] in by_text,
              "новые Heading 1/2 вошли в оглавление",
              "новые Heading 1/2 не вошли в оглавление")
        check(all(re.search(r"\t\d+$", by_text.get(text, ""))
                  for text in (spec["probe_h1"], spec["probe_h2"])),
              "у новых пунктов есть номера страниц",
              "у нового пункта оглавления нет номера страницы")
        updated_entries = set(by_text)
        check(old_entries.issubset(updated_entries),
              "старые пункты сохранены после полного обновления",
              "полное обновление потеряло старые пункты: %s"
              % sorted(old_entries - updated_entries)[:3])
        headings = set(paragraph.text.strip() for paragraph in updated.paragraphs
                       if paragraph.style.name in ("Heading 1", "Heading 2"))
        check(updated_entries == headings,
              "обновлённое оглавление совпало со всеми заголовками",
              "обновлённое оглавление разошлось с заголовками")
    finally:
        shutil.rmtree(probe_dir, ignore_errors=True)
    check(hashlib.sha256(open(path, "rb").read()).hexdigest() == source_hash,
          "регрессия не изменила исходный DOCX",
          "регрессия изменила исходный DOCX")


def verify(lang, run_toc_regression=False, verify_open_request=False):
    spec = MANUALS[lang]
    path = os.path.join(DOCS, spec["file"])
    problems = []
    log = []

    def check(condition, ok_text, bad_text):
        (log if condition else problems).append(ok_text if condition else bad_text)

    check(os.path.isfile(path), "файл найден: " + spec["file"], "нет файла: " + path)
    if not os.path.isfile(path):
        return problems, log

    doc = Document(path)
    normal = doc.styles["Normal"]
    check(normal.font.name == FONT, "основной шрифт: " + str(normal.font.name),
          "основной шрифт не Times New Roman: %s" % normal.font.name)
    check(normal.font.size == Pt(14), "основной кегль: 14 пт",
          "основной кегль не 14 пт: %s" % normal.font.size)
    check(normal.paragraph_format.first_line_indent is not None
          and abs(normal.paragraph_format.first_line_indent.cm - 1.25) < 0.01,
          "красная строка: 1,25 см",
          "красная строка не задана: %s" % normal.paragraph_format.first_line_indent)
    check(str(normal.paragraph_format.line_spacing_rule).startswith("ONE_POINT_FIVE"),
          "интервал: полуторный",
          "интервал не полуторный: %s" % normal.paragraph_format.line_spacing_rule)

    for name in ("Heading 1", "Heading 2"):
        style = doc.styles[name]
        size, font = inherited(style, "size"), inherited(style, "name")
        check(size == Pt(16), "%s: 16 пт" % name, "%s не 16 пт: %s" % (name, size))
        check(font == FONT, "%s: Times New Roman" % name,
              "%s не Times New Roman: %s" % (name, font))

    alien_font, alien_size = set(), set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name not in (None, FONT):
                alien_font.add(run.font.name)
            if run.font.size is not None and run.font.size not in (Pt(14), Pt(16), Pt(20), Pt(28)):
                alien_size.add(run.font.size.pt)
    check(not alien_font, "чужих шрифтов в тексте нет", "чужие шрифты: %s" % alien_font)
    check(not alien_size, "чужих кеглей в тексте нет", "чужие кегли: %s" % alien_size)

    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        document_xml = zf.read("word/document.xml").decode("utf-8")
        types = zf.read("[Content_Types].xml").decode("utf-8")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        settings = zf.read("word/settings.xml").decode("utf-8")
        footnotes = zf.read("word/footnotes.xml").decode("utf-8") \
            if "word/footnotes.xml" in names else ""
        footers = "".join(zf.read(name).decode("utf-8")
                          for name in names if name.startswith("word/footer"))

    check("word/footnotes.xml" in names, "часть сносок на месте", "нет части word/footnotes.xml")
    check("footnotes+xml" in types, "тип содержимого сносок объявлен",
          "в [Content_Types].xml нет типа для сносок")
    check("relationships/footnotes" in rels, "связь со сносками объявлена",
          "в связях документа нет ссылки на сноски")
    note_tags = re.findall(r'<w:footnote\b[^>]*>', footnotes)
    real_notes = [re.search(r'w:id="(-?\d+)"', tag).group(1)
                  for tag in note_tags if "w:type=" not in tag]
    refs = re.findall(r'<w:footnoteReference w:id="(-?\d+)"', document_xml)
    footnotes_match = (len(refs) > 0
                       and len(set(real_notes)) == len(real_notes)
                       and len(set(refs)) == len(refs)
                       and set(real_notes) == set(refs))
    check(footnotes_match, "сноски и ссылки совпадают: %s" % ", ".join(refs),
          "сноски и ссылки не сходятся: определения %s, ссылки %s" % (real_notes, refs))
    check(real_notes == [str(number) for number in range(1, 10)],
          "обычные идентификаторы сносок после Word: 1–9",
          "ожидались идентификаторы сносок 1–9 после Word, найдены %s" % real_notes)

    check("TOC \\o" in document_xml, "поле автооглавления вставлено", "нет поля TOC")
    check(len(re.findall(r'<w:updateFields\b[^>]*w:val="true"', settings)) == 1,
          "запрос обновления полей при открытии сохранён",
          "нет единственного w:updateFields=true после Word-финализации")
    toc_beginnings = re.findall(
        r'<w:fldChar\b(?=[^>]*w:fldCharType="begin")(?=[^>]*w:dirty="true")[^>]*/>',
        document_xml)
    check(bool(toc_beginnings), "поле TOC помечено как требующее обновления",
          "поле TOC не помечено w:dirty=true")
    check(" PAGE " in footers, "поле номера страницы стоит в колонтитуле",
          "в колонтитулах нет поля PAGE")

    entries = toc_entries(doc)
    check(len(entries) >= 10, "оглавление собрано: %d пунктов" % len(entries),
          "оглавление пустое или обрезано: %d пунктов" % len(entries))
    numbered = [paragraph.text for paragraph in entries if re.search(r"\t\d+$", paragraph.text)]
    check(len(numbered) == len(entries), "у всех пунктов оглавления есть номер страницы",
          "пунктов без номера страницы: %d" % (len(entries) - len(numbered)))
    ragged = [paragraph for paragraph in entries
              if paragraph.paragraph_format.first_line_indent is None
              or paragraph.paragraph_format.first_line_indent.cm > 0.01]
    check(not ragged, "пункты оглавления идут без красной строки",
          "пунктов оглавления с красной строкой: %d" % len(ragged))
    in_toc = set(toc_text(paragraph) for paragraph in entries)
    in_body = set(paragraph.text.strip() for paragraph in doc.paragraphs
                  if paragraph.style.name in ("Heading 1", "Heading 2"))
    check(in_toc == in_body, "оглавление совпадает с заголовками (%d)" % len(in_body),
          "оглавление разошлось с заголовками: нет %s, лишние %s"
          % (sorted(in_body - in_toc)[:3], sorted(in_toc - in_body)[:3]))
    check("<w:titlePg" in document_xml, "у первой страницы свой колонтитул (без номера)",
          "титульный лист не отделён от нумерации")

    engine_text = io.open(os.path.join(HERE, "engine.py"), encoding="utf-8").read()
    figures = len(re.findall(r'"[a-z0-9-]+"',
                             engine_text.split("FIG_ORDER = [")[1].split("]")[0]))
    images = [name for name in names if name.startswith("word/media/")]
    check(len(images) == figures, "картинок внедрено: %d" % len(images),
          "картинок внедрено %d, ожидалось %d" % (len(images), figures))

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    captions = re.findall(spec["figure"], text)
    mentions = set(int(number) for number in re.findall(spec["mention"], text, re.IGNORECASE))
    missing = [number for number in captions if int(number) not in mentions]
    check(not missing, "все %d рисунков упомянуты в тексте" % len(captions),
          "рисунки без ссылки в тексте: %s" % missing)
    check([int(number) for number in captions] == list(range(1, len(captions) + 1)),
          "нумерация рисунков сплошная", "нумерация рисунков нарушена: %s" % captions)
    review_lines = [paragraph.text for paragraph in doc.paragraphs
                    if spec["review_caption"] in paragraph.text]
    check(len(review_lines) == 1 and re.match(spec["figure"], review_lines[0])
          and int(re.match(spec["figure"], review_lines[0]).group(1)) == 25,
          "снимок Review имеет номер 25",
          "снимок Review не найден под номером 25: %s" % review_lines)

    table_caps = re.findall(spec["table"], text)
    check([int(number) for number in table_caps] == list(range(1, len(table_caps) + 1)),
          "таблиц: %d, нумерация сплошная" % len(table_caps),
          "нумерация таблиц нарушена: %s" % table_caps)
    check(len(doc.tables) == len(table_caps), "подписей к таблицам столько же, сколько таблиц",
          "таблиц %d, подписей %d" % (len(doc.tables), len(table_caps)))

    if lang == "ru":
        subprocess.check_call([sys.executable, os.path.join(HERE, "dump_loc.py")])
        loc = json.load(io.open(os.path.join(HERE, "loc_ru.json"), encoding="utf-8"))
        known = set()
        for value in loc.values():
            string = value.strip()
            known.add(string)
            known.add(string.replace("◀", "").replace("▶", "").replace("⌂", "")
                      .replace("☰", "").replace("⟳", "").strip())
            known.add(string.rstrip(". …:").strip())
            known.add(string.replace("◀", "").replace("▶", "").replace("⌂", "")
                      .replace("☰", "").replace("⟳", "").strip().rstrip(". …:").strip())
        prose = {
            ".xls", "1-3", "8-", "Microsoft Print to PDF", "[BASENAME]_часть_[FILENUMBER###]",
            "Блокнот", "Договор_часть_001", "Договор_часть_002", "взяв рукой",
            "каждая страница отдельным файлом", "плывут", "поедут", "склеивания",
            "том на триста страниц, смотрите раздел четыре", "файл повреждён", "Пуск",
            "Хорошо", "Нормально", "Инструкция по работе с программой: открыть",
            "Настройках", "страница целиком картинкой", "только текста", "у вас последняя",
            "Обновить таблицу", "Обновить целиком", "Обновить только номера страниц",
            "Заголовок 1", "Заголовок 2", "Исходный вид", "Поменять местами",
            "Предыдущее", "Сравнение текста", "съехали", "Красный — удалено",
            "Зелёный — добавлено", "Что нового", "сторонних компонентах",
        }
        lower = set(value.lower() for value in known)

        def is_caption(quote):
            if quote in known or quote.lower() in lower:
                return True
            return any(value.strip().startswith(quote) or ("«%s»" % quote) in value.strip()
                       for value in loc.values())

        unknown = [quote for quote in sorted(set(re.findall("«([^»]{2,60})»", text)))
                   if not is_caption(quote) and quote not in prose]
        check(not unknown, "все подписи в кавычках есть в каталоге программы",
              "подписи, которых нет в программе: %s" % ", ".join(unknown))

    banned_fingerprints = {
        "610a4b77c77b3ffd", "2c0cda61e536f2c5", "fe4e39d2d1e3c6b5",
        "2821bc2655a8343f", "d74a2fa10ca4be11", "f3dcd5e6ce37de3e",
        "dab70bc2d883a287",
    }

    def fingerprint(word):
        return hashlib.sha256(word.encode("utf-8")).hexdigest()[:16]

    found = []
    for raw in re.findall(r"[^\W\d_]{3,}", text.lower(), re.UNICODE):
        for size in range(3, min(len(raw), 14) + 1):
            if fingerprint(raw[:size]) in banned_fingerprints:
                found.append(raw)
                break
    check(not found, "рабочих упоминаний нет",
          "найдены рабочие упоминания: %s" % sorted(set(found)))

    if run_toc_regression:
        verify_toc_update(path, spec, check, verify_open_request=verify_open_request)

    print("--- %s: %s" % (lang, spec["file"]))
    print("Заголовков: %d" % len([paragraph for paragraph in doc.paragraphs
                                  if paragraph.style.name.startswith("Heading")]))
    for line in log:
        print("  ок    ", line)
    for line in problems:
        print("  ПЛОХО ", line)
    print("итог:", "ВСЁ СОШЛОСЬ" if not problems else "ЕСТЬ ЗАМЕЧАНИЯ (%d)" % len(problems))
    print()
    return problems, log


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--lang", choices=("ru", "en"))
    parser.add_argument("--both", action="store_true")
    parser.add_argument("--toc-regression", action="store_true")
    parser.add_argument("--verify-open-request", action="store_true")
    args = parser.parse_args()
    languages = ("ru", "en") if args.both or not args.lang else (args.lang,)
    failed = False
    for language in languages:
        problems, _ = verify(language, args.toc_regression, args.verify_open_request)
        failed = failed or bool(problems)
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
